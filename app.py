# app.py
from flask import Flask, render_template, session, request, redirect, url_for, abort, flash, make_response, \
   render_template_string, send_file, get_flashed_messages
from flask_migrate import Migrate
from flask_login import LoginManager, login_required, login_user, current_user, logout_user
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from datetime import datetime
from dotenv import load_dotenv
import os
import re
import io

import html
from bs4 import BeautifulSoup, NavigableString
import bleach

# LLM SDKs
import google.generativeai as genai
import openai
from openai import OpenAI


# Documents / PDF
import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert

from fpdf import FPDF
from weasyprint import HTML
import pdfkit
from PyPDF2 import PdfReader
import PyPDF2


# Parsing / templating
import markdown
from bs4 import BeautifulSoup
from markupsafe import Markup


# Images
from PIL import Image


# DB models
from models import db, Conversation
from models import User, Resume, ResumeInput, UsageLog, CoverLetter

# HTTP utilities
from io import BytesIO
import requests


# Flask-Babel
from flask_babel import Babel, _


# Load env
load_dotenv()


# Flask app
app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'your-secret-key')
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv("DATABASE_URL").strip("'").strip('"')
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024  # 5MB
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False


# LLM keys
openai.api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
GEMINI_KEY = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
if GEMINI_KEY:
   genai.configure(api_key=GEMINI_KEY)


# Languages
app.config['LANGUAGES'] = {
   'en': 'English', 'es': 'Spanish', 'fr': 'French', 'de': 'German',
   'zh': 'Chinese', 'ja': 'Japanese', 'ru': 'Russian', 'ar': 'Arabic',
   'hi': 'Hindi', 'pt': 'Portuguese', 'it': 'Italian', 'ko': 'Korean',
}

def get_locale():
   return request.accept_languages.best_match(app.config['LANGUAGES'].keys())


babel = Babel(app, locale_selector=get_locale)


# Folders
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)


# DB / Migrate
db.init_app(app)
migrate = Migrate(app, db)


# Login
login_manager = LoginManager(app)
login_manager.login_view = 'login'


@login_manager.user_loader
def load_user(user_id):
   return db.session.get(User, int(user_id))


# -----------------------------
# Routes / Utilities
# -----------------------------
@app.route("/")
def home():
   current_year = datetime.now().year
   return render_template("home.html", current_year=current_year)


@app.route("/admin")
def admin_dashboard():
   resume_count = Resume.query.count()
   generate_count = UsageLog.query.filter_by(action="generate_resume").count()
   return render_template("admin_dashboard.html", resume_count=resume_count, generate_count=generate_count)


@app.route("/dashboard")
@login_required
def dashboard():
   return render_template("dashboard.html")


ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'rtf'}


def allowed_file(filename):
   return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def allowed_photo(filename):
   return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'png', 'jpg', 'jpeg'}


def make_all_links_clickable(text: str) -> str:
   """Convert Markdown-style links + plain URLs to clickable anchors."""
   text = re.sub(
       r'\[([^\]]+)\]\(((?:https?://)[^\s)]+)\)',
       r'<a href="\2" target="_blank" rel="noopener noreferrer">\1</a>',
       text
   )
   text = re.sub(
       r'(?<!["\'>=])((https?://[^\s<>"\'\)]+))',
       r'<a href="\1" target="_blank" rel="noopener noreferrer">\1</a>',
       text
   )
   return text


def add_hyperlink(paragraph, text, url):
    """
    Insert a clickable hyperlink into a python-docx paragraph.
    Returns the created <w:hyperlink> element.
    """
    # Creating the relationship in the document part
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")  # Word default link blue
    rPr.append(color)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def extract_text(file_path, ext):
   if ext == 'pdf':
       text = ""
       with open(file_path, 'rb') as f:
           reader = PdfReader(f)
           for page in reader.pages:
               page_text = page.extract_text()
               if page_text:
                   text += page_text + "\n"
   elif ext == 'docx':
       d = docx.Document(file_path)
       text = "\n".join([para.text for para in d.paragraphs])
   elif ext == 'txt':
       with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
           text = f.read()
   elif ext == 'rtf':
       from striprtf.striprtf import rtf_to_text
       with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
           text = rtf_to_text(f.read())
   else:
       raise ValueError("Unsupported file type")
   clean_text = text.replace('\x00', '')
   return make_all_links_clickable(clean_text)


@app.route("/upload", methods=["GET", "POST"])
@login_required
def upload_resume():
   VALID_MODELS = ['gpt-4o-mini', 'gpt-4.1-mini', 'gemini-2.5-flash']
   if request.method == "POST":
       if 'resume' not in request.files:
           return "No file part", 400
       file = request.files['resume']
       if file.filename == '':
           return "No selected file", 400


       upload_folder = app.config['UPLOAD_FOLDER']
       os.makedirs(upload_folder, exist_ok=True)


       # Optional photo
       photo_path = None
       if 'photo' in request.files:
           photo = request.files['photo']
           if photo.filename != '' and allowed_photo(photo.filename):
               photo_filename = secure_filename(photo.filename)
               photo_path = os.path.join(upload_folder, photo_filename)
               try:
                   photo.save(photo_path)
               except Exception as e:
                   return f"Photo save error: {str(e)}", 500


       selected_template = request.form.get("template", "professional")
       selected_model = request.form.get("model", "gpt-4o-mini")
       if selected_model not in VALID_MODELS:
           selected_model = "gpt-4o-mini"


       if file and allowed_file(file.filename):
           filename = secure_filename(file.filename)
           ext = filename.rsplit('.', 1)[1].lower()
           path = os.path.join(upload_folder, filename)
           try:
               file.save(path)
               content = extract_text(path, ext)


               # AI Resume Generation
               professional_resume = generate_resume_text_unified(content, model=selected_model)


               enhanced_resume = Resume(
                   user_id=current_user.id,
                   type="ai_enhanced",
                   original_name=f"Enhanced_{filename}",
                   text_content=professional_resume,
                   ai_feedback=f"Enhanced using {selected_model}",
                   photo_path=photo_path,
                   template_name=selected_template,
                   created_at=datetime.now()
               )
               original_resume = Resume(
                   user_id=current_user.id,
                   type="uploaded_original",
                   original_name=filename,
                   text_content=content,
                   created_at=datetime.now()
               )


               db.session.add(original_resume)
               db.session.add(enhanced_resume)
               db.session.commit()


               return redirect(url_for('view_enhanced_resume', id=enhanced_resume.id))


           except Exception as e:
               if photo_path and os.path.exists(photo_path):
                   os.remove(photo_path)
               if path and os.path.exists(path):
                   os.remove(path)
               return f"Error processing file: {str(e)}", 500


       return "Invalid file type. Supported formats: PDF, DOCX, TXT, RTF", 400


   return render_template("upload.html", models=VALID_MODELS)


# ----- Unified AI Resume Generator -----
def generate_resume_text_unified(content: str, model: str = "gpt-4o-mini") -> str:
   STRICT_PROMPT = f"""
You are a professional resume editing assistant.


Enhance the following resume content so it reads clearly, professionally, and matches modern ATS-friendly formatting.


IMPORTANT RULES:
1. DO NOT invent, fabricate, or add any new achievements, dates, roles, tools, or details not present in the provided text.
2. Only rephrase, restructure, and improve grammar/design while keeping all facts, dates, and details exactly the same.
3. Preserve the existing section structure and order (e.g. Professional Summary, Skills, Education, Projects, Experience, Certifications, Languages).
4. Keep formatting clean and ATS-readable (no special characters, no tables).
5. Output ONLY the improved resume content — no advice, no explanations.


RESUME CONTENT TO IMPROVE:
{content}
"""
   try:
       if model.startswith("gemini"):
           if not GEMINI_KEY:
               raise RuntimeError("GEMINI_API_KEY is not set in environment.")
           gm = genai.GenerativeModel(model)
           resp = gm.generate_content(STRICT_PROMPT)
           output_text = getattr(resp, "text", str(resp)).strip()
       else:
           response = client.chat.completions.create(
               model=model,
               messages=[
                   {"role": "system", "content": "You are a professional resume writer. Follow the rules strictly."},
                   {"role": "user", "content": STRICT_PROMPT}
               ],
               temperature=0,
               max_tokens=2000
           )
           output_text = response.choices[0].message.content.strip()


       # block suspicious added sections
       suspicious_sections = ["Awards", "Volunteer", "Extra Projects", "Interests"]
       for section in suspicious_sections:
           if section.lower() in output_text.lower() and section.lower() not in content.lower():
               output_text = re.sub(rf"{section}.*", "", output_text, flags=re.IGNORECASE)


       return output_text
   except Exception as e:
       return f"Error generating resume with model '{model}': {e}"


@app.route("/resumes")
@login_required
def resume_list():
   resumes = Resume.query.filter_by(user_id=current_user.id).order_by(Resume.created_at.desc()).all()
   return render_template("resumes_list.html", resumes=resumes)


def generate_professional_resume(original_content, model="gpt-4o-mini"):
   try:
       prompt = f"Here's my resume, please enhance it:\n{original_content}"
       response = client.chat.completions.create(
           model=model,
           messages=[
               {"role": "system",
                "content": ("You are an expert resume writer. "
                            "Rewrite and enhance the following resume content to be more clear, impactful, and professional. "
                            "Do NOT add new sections, comments, or explanations. "
                            "Return ONLY the improved resume content, preserving structure as much as possible.")},
               {"role": "user", "content": prompt}
           ],
           temperature=0.3,
           max_tokens=2000
       )
       return response.choices[0].message.content.strip()
   except Exception as e:
       return f"⚠️ AI Enhancement Error: {str(e)}"


def replace_placeholder_links(md_text: str, links: dict) -> str:
   pattern = re.compile(r'\[([^\]]+)\]\((?:#|)\)')
   def repl(m):
       label = m.group(1)
       return f'[{label}]({links.get(label, "#")})'
   md_text = pattern.sub(repl, md_text)
   bare_pattern = re.compile(r'\[([^\]]+)\](?!\()')
   def bare_repl(m):
       label = m.group(1)
       return f'[{label}]({links.get(label, "#")})'
   md_text = bare_pattern.sub(bare_repl, md_text)
   return md_text


@app.route("/resume/enhanced/<int:id>")
@login_required
def view_enhanced_resume(id):
    resume = Resume.query.get_or_404(id)
    if resume.user_id != current_user.id:
        abort(403)

    # Markdown → HTML + link fixes
    md = replace_placeholder_links(resume.text_content, LINK_MAP)
    html = markdown.markdown(md, extensions=['extra', 'sane_lists'])
    html = make_all_links_clickable(html)

    # Normalize anchors
    soup = BeautifulSoup(html, "html.parser")
    for a in soup.find_all('a', href=True):
        href = a['href'].strip()
        if href.startswith('www.'):
            a['href'] = f'https://{href}'
        a['target'] = '_blank'
        a['rel'] = 'noopener noreferrer'
    html = str(soup)

    return render_template(
        "enhanced_resume_view.html",
        resume=resume,
        html_content=html,
        resume_markdown=resume.text_content,
        selected_template=(resume.template_name or "professional-blue")
    )


@app.route('/_wrap_preview', methods=['POST'])
@login_required
def wrap_preview():
    data = request.get_json() or {}

    md = replace_placeholder_links(data.get('markdown') or '', LINK_MAP)
    html = markdown.markdown(md, extensions=['extra', 'sane_lists'])
    html = make_all_links_clickable(html)                    # http(s), www, emails
    html = ensure_label_links_in_html(html, LINK_MAP)        # "LinkedIn" -> <a ...>LinkedIn</a>

    soup = BeautifulSoup(html, "html.parser")
    for a in soup.find_all('a', href=True):
        href = a['href'].strip()
        if href.startswith('www.'):
            a['href'] = f'https://{href}'
        a['target'] = '_blank'
        a['rel'] = 'noopener noreferrer'
    return str(soup)


@app.route("/resume/compare/<int:original_id>/<int:enhanced_id>")
@login_required
def compare_resumes(original_id, enhanced_id):
   original = Resume.query.get_or_404(original_id)
   enhanced = Resume.query.get_or_404(enhanced_id)
   if original.user_id != current_user.id or enhanced.user_id != current_user.id:
       abort(403)
   return render_template("compare_resumes.html", original=original, enhanced=enhanced)


@app.route("/resume/<int:id>/delete", methods=["POST"], endpoint='delete_resume_endpoint')
@login_required
def delete_resume(id):
   resume = Resume.query.get_or_404(id)
   if resume.user_id != current_user.id:
       abort(403)
   db.session.delete(resume)
   db.session.commit()
   return redirect(url_for('my_resumes'))

@app.route("/resumes/bulk-delete", methods=["POST"])
@login_required
def bulk_delete_resumes():
    ids = request.form.getlist("resume_ids")
    if not ids:
        flash("Please select at least one resume to delete.", "warning")
        return redirect(url_for('my_resumes'))

    try:
        ids = [int(i) for i in ids]
    except ValueError:
        abort(400)

    to_delete = Resume.query.filter(
        Resume.user_id == current_user.id,
        Resume.id.in_(ids)
    ).all()

    count = len(to_delete)
    for r in to_delete:
        db.session.delete(r)
    db.session.commit()

    flash(f"Deleted {count} resume(s).", "success")
    return redirect(url_for('my_resumes'))



@app.route("/resume/<int:id>", endpoint='view_resume_endpoint')
@login_required
def view_resume(id):
   resume = Resume.query.get_or_404(id)
   if resume.user_id != current_user.id:
       abort(403)
   return render_template("resume_view.html", resume=resume)


@app.route("/resume/<int:id>/edit", methods=["GET", "POST"])
@login_required
def edit_resume(id):
    resume = Resume.query.get_or_404(id)
    if resume.user_id != current_user.id:
        abort(403)

    if request.method == "POST":
        new_text = (request.form.get("text_content") or "").strip()
        new_template = (request.form.get("template_name") or "").strip()

        if not new_text:
            flash("Resume content cannot be empty", "error")
            return render_template("edit_resume.html", resume=resume)

        resume.text_content = new_text
        if new_template:
            resume.template_name = new_template

        db.session.commit()
        flash("Resume updated successfully", "success")
        # 👉 send them to the enhanced preview page
        return redirect(url_for('view_enhanced_resume', id=resume.id))

    return render_template("edit_resume.html", resume=resume)



@app.route("/register", methods=["GET", "POST"])
def register():
   if request.method == "POST":
       username = request.form["username"]
       email = request.form["email"]
       password = request.form["password"]
       confirm_password = request.form["confirm_password"]
       if password != confirm_password:
           return render_template("register.html", error="Passwords do not match")
       existing_user = User.query.filter_by(email=email).first()
       if existing_user:
           return render_template("register.html", error="Email is already registered")
       hashed_password = generate_password_hash(password)
       user = User(name=username, email=email, password=hashed_password)
       db.session.add(user)
       db.session.commit()
       flash("Registration successful! Please log in.", "success")
       return redirect(url_for("login"))
   return render_template("register.html")


@app.route("/login", methods=["GET", "POST"])
def login():
   if request.method == "POST":
       user = User.query.filter_by(email=request.form['email']).first()
       if user and check_password_hash(user.password, request.form['password']):
           login_user(user)
           return redirect(url_for('home'))
   return render_template("login.html")


@app.route('/logout')
@login_required
def logout():
   logout_user()
   return redirect(url_for('home'))

def ensure_period(text: str) -> str:
   return text.strip() + ("" if text.endswith((".", "!", "?")) else ".")


# Career Advice (tokens/speed removed)
VALID_MODELS = ['gpt-4o-mini', 'gpt-4.1-mini', 'gemini-2.5-flash']
_GEMINI_KEY = GEMINI_KEY


def call_gemini_flash_api(question, history=None, max_turns=10):
   """
   Uses Gemini 2.5 Flash with preserved chat context.
   Returns: advice_text (str)
   """
   if not _GEMINI_KEY:
       raise RuntimeError("Gemini API key not set. Provide GEMINI_API_KEY or GOOGLE_API_KEY.")

   gm = genai.GenerativeModel("gemini-2.5-flash")

   hist = []
   if history:
       for turn in list(history)[-max_turns:]:
           role = "user" if getattr(turn, "role", "user") == "user" else "model"
           content = getattr(turn, "content", "")
           if content:
               hist.append({"role": role, "parts": [content]})

   chat = gm.start_chat(history=hist)
   gemini_resp = chat.send_message(question.strip())

   advice = (getattr(gemini_resp, "text", "") or str(gemini_resp)).strip()
   advice = advice.rstrip('.!,;、。 ')
   if not advice.endswith('?'):
       advice += '?'

   return advice


@app.route('/career-advice', methods=['GET', 'POST'])
@login_required
def career_advice():
   image_url = None
   if request.method == 'POST':
       question = request.form.get('question', '').strip()
       selected_model = request.form.get('model', 'gpt-4o-mini')
       language = request.form.get('language', 'en').strip()

       if selected_model not in VALID_MODELS:
           flash("Invalid model selected.", "error")
           return redirect(url_for('career_advice'))

       DEFAULT_MARKERS = {"your question:", "...", "n/a", "?"}
       if not question or question.lower() in DEFAULT_MARKERS:
           flash("Please ask a more specific question so I can help you better.", "error")
           return redirect(url_for('career_advice'))

       # optional file attachment
       file_text = ""
       attachment = request.files.get('attachment')
       if attachment and attachment.filename:
           filename = attachment.filename.lower()
           try:
               content_bytes = attachment.read()
               if filename.endswith('.txt'):
                   file_text = content_bytes.decode('utf-8', errors='ignore')
               elif filename.endswith('.pdf'):
                   reader = PyPDF2.PdfReader(io.BytesIO(content_bytes))
                   file_text = "\n".join(page.extract_text() or "" for page in reader.pages)
               elif filename.endswith('.docx'):
                   d = docx.Document(io.BytesIO(content_bytes))
                   file_text = "\n".join(para.text for para in d.paragraphs)
               else:
                   flash("Unsupported file type. Please upload TXT, PDF, or DOCX files.", "warning")
           except Exception as e:
               flash(f"Failed to process attachment: {str(e)}", "error")
               return redirect(url_for('career_advice'))

       user_input = question
       if file_text:
           user_input += f"\n\nAdditional context from the uploaded file:\n{file_text}"

       # history
       history = Conversation.query.filter_by(user_id=current_user.id)\
                                   .order_by(Conversation.timestamp).all()

       try:
           if selected_model == "gemini-2.5-flash":
               advice = call_gemini_flash_api(user_input, history=history, max_turns=10)
               advice = strip_markdown(advice)
               advice = ensure_ends_with_question(advice)

           else:
               messages = [{
                   "role": "system",
                   "content": (
                       "You are an expert career advisor. "
                       "You remember the user's previous questions and your previous answers naturally. "
                       "Do not mention memory limitations."
                   )
               }]
               for turn in history[-20:]:
                   messages.append({"role": turn.role, "content": turn.content})
               messages.append({"role": "user", "content": user_input})

               response = client.chat.completions.create(
                   model=selected_model,
                   messages=messages,
                   max_tokens=1200,
                   temperature=0.7
               )

               advice = response.choices[0].message.content.strip()
               advice = strip_markdown(advice)
               advice = ensure_ends_with_question(advice)

           # Save convo
           db.session.add(Conversation(
               user_id=current_user.id,
               role='user',
               content=user_input,
               timestamp=datetime.now()
           ))
           db.session.add(Conversation(
               user_id=current_user.id,
               role='assistant',
               content=advice,
               timestamp=datetime.now()
           ))
           db.session.commit()

       except Exception as e:
           flash(f"An error occurred: {str(e)}", "error")
           return redirect(url_for('career_advice'))

       history = Conversation.query.filter_by(user_id=current_user.id)\
                                   .order_by(Conversation.timestamp).all()

       return render_template(
           "career_advice.html",
           advice=advice,
           image_url=image_url,
           history=history,
           selected_model=selected_model,
           selected_language=language,
           last_question=question,
           question=""
       )

   # GET
   history = Conversation.query.filter_by(user_id=current_user.id)\
                               .order_by(Conversation.timestamp).all()
   return render_template(
       "career_advice.html",
       advice=None,
       image_url=None,
       history=history,
       selected_model="gpt-4o-mini",
       selected_language="en",
       question=""
   )

# Language helpers (place ABOVE routes)
try:
    from langdetect import detect as _ld_detect  # pip install langdetect
except Exception:
    _ld_detect = None

def _detect_lang(text: str, default="en") -> str:
    """
    Detect a language code from freeform text.
    Falls back to 'default' if detection fails.
    """
    text = (text or "").strip()
    if not text:
        return default
    if _ld_detect:
        try:
            lang = _ld_detect(text)
            # normalize variants
            alias = {"zh-cn": "zh", "zh-tw": "zh"}
            return alias.get(lang.lower(), lang.lower())
        except Exception:
            pass

    # Tiny keyword heuristics as a fallback
    t = text.lower()
    if any(w in t for w in ["bewerbung", "kenntnisse", "erfahrung", "sehr geehrte", "mit freundlichen grüßen"]):
        return "de"
    if any(w in t for w in ["estimado", "experiencia", "habilidades", "atentamente"]):
        return "es"
    if any(w in t for w in ["monsieur", "madame", "compétences", "poste", "cordialement"]):
        return "fr"
    if any(w in t for w in ["عزيزي", "السيد", "شركة", "خبرة", "مع فائق الاحترام"]):
        return "ar"
    return default

# Minimal localization dictionary for labels + phrasing
_LOCALE = {
    "de": {
        "email": "E-Mail",
        "phone": "Telefon",
        "date": "Datum",
        "recipient_block": "An die Personalabteilung\n[Unternehmensname]\n[Unternehmensadresse]\n[Stadt, Land]",
        "salutation": "Sehr geehrte Damen und Herren,",
        "closing": "Mit freundlichen Grüßen,",
        "instructions": (
            "Schreibe den gesamten Brief ausschließlich auf Deutsch. "
            "Verwende einen förmlichen Ton (Sie-Form). "
            "KEINE englischen Wörter oder Sätze (außer Eigennamen)."
        ),
        "date_fmt": "%d.%m.%Y",
    },
    "en": {
        "email": "Email",
        "phone": "Phone",
        "date": "Date",
        "recipient_block": "Hiring Manager\n[Company Name]\n[Company Address]\n[City, Country]",
        "salutation": "Dear Hiring Manager,",
        "closing": "Sincerely,",
        "instructions": (
            "Write the entire letter only in English. "
            "No other languages (except proper nouns)."
        ),
        "date_fmt": "%d %B %Y",
    },
    "fr": {
        "email": "E-mail",
        "phone": "Téléphone",
        "date": "Date",
        "recipient_block": "Service Recrutement\n[Nom de l’entreprise]\n[Adresse de l’entreprise]\n[Ville, Pays]",
        "salutation": "Madame, Monsieur,",
        "closing": "Cordialement,",
        "instructions": (
            "Rédigez la lettre entièrement en français, ton professionnel, vouvoiement. "
            "Aucune autre langue (sauf noms propres)."
        ),
        "date_fmt": "%d %B %Y",
    },
    "es": {
        "email": "Correo electrónico",
        "phone": "Teléfono",
        "date": "Fecha",
        "recipient_block": "Departamento de RR. HH.\n[Nombre de la empresa]\n[Dirección de la empresa]\n[Ciudad, País]",
        "salutation": "Estimado/a responsable de selección:",
        "closing": "Atentamente,",
        "instructions": (
            "Escribe toda la carta únicamente en español, tono formal. "
            "No utilices otros idiomas (salvo nombres propios)."
        ),
        "date_fmt": "%d de %B de %Y",
    },
    "ar": {
        "email": "البريد الإلكتروني",
        "phone": "الهاتف",
        "date": "التاريخ",
        "recipient_block": "قسم شؤون الموظفين\n[اسم الشركة]\n[عنوان الشركة]\n[المدينة، البلد]",
        "salutation": "السيد/السيدة المحترم(ة)،",
        "closing": "مع فائق الاحترام،",
        "instructions": (
            "اكتب الرسالة كاملة باللغة العربية الفصحى وبأسلوب رسمي. "
            "لا تستخدم أي لغات أخرى (باستثناء الأعلام)."
        ),
        "date_fmt": "%d %B %Y",
    },
}

def _locale_for(lang: str):
    """
    Return localized labels + rules for the language code.
    Falls back to English if missing.
    """
    lang = (lang or "en").lower()
    return _LOCALE.get(lang, _LOCALE["en"])


# -----------------------------
# Global link map + styles
# -----------------------------
LINK_MAP = {
    "LinkedIn": "https://www.linkedin.com/in/dennis-charles53/",
    "GitHub": "https://github.com/Dennis2y",
    "Movie Project": "https://github.com/Dennis2y/Dennis2y-Movie-Project-SQL-HTML-API",
    "Best-Buy-2.0": "https://github.com/Dennis2y/Best-Buy-2.0",
    "Credentials": "https://drive.google.com/file/d/1Hav3gSC4TMhOz091ZosUlcVvXJOTwRRi/view?pli=1",
    "Portfolio": "https://dennischarles.dev",
    "Resume": "https://example.com/dennis_resume.pdf",
    "View Credentials": "https://drive.google.com/file/d/1Hav3gSC4TMhOz091ZosUlcVvXJOTwRRi/view?pli=1",
    "Link": "https://github.com/Dennis2y/Dennis2y-Movie-Project-SQL-HTML-API",
}

ANCHOR_CSS = """
a {
    color: #2563eb;           /* blue */
    text-decoration: underline;
}
a::after {
    content: "";
    font-size: 0.9em;
}
"""

# Cover Letter route
@app.route("/cover-letter", methods=["GET", "POST"])
@login_required
def cover_letter():
    VALID_MODELS = ['gpt-4o-mini', 'gpt-4.1-mini']
    selected_model = "gpt-4o-mini"

    if request.method == "POST":
        selected_model = request.form.get("model", "gpt-4o-mini")
        if selected_model not in VALID_MODELS:
            selected_model = "gpt-4o-mini"

        # Inputs
        address = request.form.get("address", "").strip()
        email = request.form.get("email", "").strip()
        phone = request.form.get("phone", "").strip()
        resume_text = request.form.get("resume_text", "").strip()
        job_info = request.form.get("job_info", "").strip()
        # Optional language hint from UI (e.g., a <select name="language">)
        lang_hint = (request.form.get("language") or "").strip().lower()

        # Basic validation
        if not address or not email or not resume_text or not job_info:
            flash("Please provide your address, email, resume text, and job information", "error")
            return render_template(
                "cover_letter.html",
                letter="",
                letter_markdown="",
                address=address,
                email=email,
                phone=phone,
                selected_model=selected_model
            )

        # Decide target language:
        lang = lang_hint or _detect_lang(job_info or resume_text or address, default="en")
        loc = _locale_for(lang)

        # Localized date
        try:
            date_today = datetime.now().strftime(loc["date_fmt"])
        except Exception:
            date_today = datetime.now().strftime("%d %B %Y")

        # Header block with localized labels
        header_block = (
            f"{address}\n"
            f"{loc['email']}: {email}\n"
            f"{loc['phone']}: {phone}\n"
            f"{loc['date']}: {date_today}\n"
        )

        # Build a VERY explicit prompt in the chosen language
        # (We still write the meta-instructions in English to keep them precise for the model,
        # but we *force* output language and forbid mixing.)
        prompt = f"""
You are an expert cover-letter writer.

OUTPUT LANGUAGE: {lang}
HARD CONSTRAINTS:
- The ENTIRE letter MUST be written ONLY in the output language above ({lang}). NO mixing with other languages (except proper nouns like company/person names).
- Use a formal business tone matching the locale.
- Use Markdown (no code fences, no tables). Do not add commentary.

FORMAT (follow EXACTLY):
1) Print this header block FIRST, UNCHANGED, keeping line breaks:
{header_block}
2) Then ONE blank line.
3) Recipient block (localized), on separate lines (no bullets):
{loc['recipient_block']}
4) Then ONE blank line.
5) Salutation line (localized):
{loc['salutation']}
6) Then ONE blank line.
7) Body: opening paragraph stating position + motivation; 1–2 paragraphs mapping skills/achievements to the job;
   closing paragraph with a polite call to action. Keep under 400 words total.
8) Closing line (localized), then on the next line the applicant's name (only as a signature):
{loc['closing']}
Dennis Charles

ADDITIONAL LOCALIZED INSTRUCTIONS:
{loc['instructions']}

JOB DETAILS (verbatim from user):
{job_info}

APPLICANT RESUME SUMMARY (verbatim from user):
{resume_text}
""".strip()

        try:
            # Use your existing text generation function for consistency
            letter_markdown = generate_resume_text_unified(prompt, model=selected_model)

            # Clean & render
            letter_markdown_clean = strip_markdown_code_fences(letter_markdown)
            letter_html = markdown.markdown(letter_markdown_clean)

            return render_template(
                "cover_letter.html",
                letter=letter_html,
                letter_markdown=letter_markdown_clean,
                address=address,
                email=email,
                phone=phone,
                selected_model=selected_model
            )
        except Exception as e:
            flash(f"Error generating cover letter: {str(e)}", "error")
            return render_template(
                "cover_letter.html",
                letter="",
                letter_markdown="",
                address=address,
                email=email,
                phone=phone,
                selected_model=selected_model
            )

    # GET
    return render_template(
        "cover_letter.html",
        letter="",
        letter_markdown="",
        address="",
        email=current_user.email,
        phone="",
        selected_model=selected_model
    )


def strip_markdown_code_fences(text):
   return re.sub(r'```[\s\S]*?```', '', text)


def remove_name_from_header(markdown_text, user_name):
   lines = markdown_text.strip().splitlines()
   if lines and user_name.lower() in lines[0].lower():
       lines = lines[1:]
   return "\n".join(lines)


def clean_letter_markdown(letter_markdown, user_email):
   email_pattern = re.compile(re.escape(user_email.strip()), re.IGNORECASE)
   lines = letter_markdown.split('\n')
   cleaned = []
   email_seen = False
   for line in lines:
       if email_pattern.search(line.strip()):
           if not email_seen:
               cleaned.append(line)
               email_seen = True
       else:
           cleaned.append(line)
   return '\n'.join(cleaned)


@app.errorhandler(404)
def page_not_found(e):
   return render_template("404.html"), 404


@app.route("/my_resumes")
@login_required
def my_resumes():
   resumes = Resume.query.filter_by(user_id=current_user.id).all()
   return render_template("resumes_list.html", resumes=resumes)


def ensure_ends_with_question(advice):
   advice = advice.strip()
   advice = advice.rstrip('.!,;、。 ')
   if not advice.endswith('?'):
       advice += '?'
   return advice


def replace_placeholder_links(text, link_mapping):
   for label, url in link_mapping.items():
       pattern = rf'\[{re.escape(label)}\]\((#|)\)'
       replacement = f'[{label}]({url})'
       text = re.sub(pattern, replacement, text)
   return text


def rewrite_markdown_links_to_matching_text(md_text):
   pattern = r'\[([^\]]+)\]\((https?://[^\)]+)\)'
   return re.sub(pattern, lambda m: f"[{m.group(2)}]({m.group(2)})", md_text)


def strip_markdown(text):
   text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
   text = re.sub(r'\*(.*?)\*', r'\1', text)
   text = re.sub(r'`(.*?)`', r'\1', text)
   text = re.sub(r'__([^_]+)__', r'\1', text)
   text = re.sub(r'^\s*[\-\*]\s+', '', text, flags=re.MULTILINE)
   text = re.sub(r"```.*?```", "", text, flags=re.DOTALL)
   return text.strip()


# --- route: replace your existing /download_template_pdf with this ---

@app.route('/download-template-pdf', methods=['POST'])
@login_required
def download_template_pdf():
    data = request.get_json() or {}
    doc_type = (data.get('docType') or 'resume').strip()
    is_markdown = bool(data.get('isMarkdown', True))
    content = (data.get('content') or "").strip()

    resume_id = data.get('resumeId')
    if not content and resume_id:
        r = Resume.query.get(int(resume_id))
        if r and r.user_id == current_user.id:
            content = (r.text_content or "").strip()
    if not content:
        content = "_(No content provided)_"

    # Markdown → HTML
    content = replace_placeholder_links(content, LINK_MAP)
    content_html = markdown.markdown(content, extensions=['extra', 'sane_lists']) if is_markdown else content

    # Make links clickable and convert label words to anchors
    content_html = make_all_links_clickable(content_html)
    content_html = ensure_label_links_in_html(content_html, LINK_MAP)

    # Normalize anchors
    soup = BeautifulSoup(content_html, "html.parser")
    for a in soup.find_all('a', href=True):
        href = (a['href'] or '').strip()
        if href.startswith('www.'):
            a['href'] = f'https://{href}'
        a['target'] = '_blank'
        a['rel'] = 'noopener noreferrer'
    content_html = str(soup)

    # CSS that WeasyPrint will honor -> blue + underline + clickable
    base_css = """
      html, body { margin:0; padding:0; }
      body { font-family: Arial, sans-serif; background:#fff; color:#111; }
      .page { margin: 40px; }
      h1{ font-size: 26px; } h2{ font-size: 18px; } h3{ font-size: 15px; }
      p, li { font-size: 13.5px; line-height: 1.55; }
      a, a:link, a:visited { color:#2563eb; text-decoration: underline; }
      @page { size: A4; margin: 18mm 16mm; }
    """

    html_doc = f"""<!doctype html>
<html>
<head>
<meta charset="utf-8">
<style>{base_css}</style>
</head>
<body>
  <div class="page">
    {content_html}
  </div>
</body>
</html>"""

    pdf = HTML(string=html_doc, base_url=request.host_url).write_pdf()
    resp = make_response(pdf)
    resp.headers['Content-Type'] = 'application/pdf'
    resp.headers['Content-Disposition'] = f'attachment; filename={doc_type.capitalize()}.pdf'
    return resp


# Global link map (used to turn [Label]() into real links)
LINK_MAP = {
    "LinkedIn": "https://www.linkedin.com/in/dennis-charles53/",
    "GitHub": "https://github.com/Dennis2y",
    "Movie Project": "https://github.com/Dennis2y/Dennis2y-Movie-Project-SQL-HTML-API",
    "Best-Buy-2.0": "https://github.com/Dennis2y/Best-Buy-2.0",
    "Credentials": "https://drive.google.com/file/d/1Hav3gSC4TMhOz091ZosUlcVvXJOTwRRi/view?pli=1",
    "Portfolio": "https://dennischarles.dev",
    "Resume": "https://example.com/dennis_resume.pdf",
    "View Credentials": "https://drive.google.com/file/d/1Hav3gSC4TMhOz091ZosUlcVvXJOTwRRi/view?pli=1",
    "Link": "https://github.com/Dennis2y/Dennis2y-Movie-Project-SQL-HTML-API",
}
link_map = LINK_MAP



def ensure_label_links_in_html(content_html: str, link_map: dict) -> str:
    """
    Finds plain text label words like 'LinkedIn' / 'GitHub' that are NOT already
    inside an <a> tag and converts them into <a href="...">Label</a>.
    This runs on HTML (after your markdown conversion).
    """
    soup = BeautifulSoup(content_html, "html.parser")

    # For each label, replace text nodes that contain the exact word (word-boundary).
    for label, url in link_map.items():
        # find all text nodes containing the label (case-sensitive by default)
        matches = soup.find_all(string=re.compile(rf'\b{re.escape(label)}\b'))
        for node in matches:
            # skip if this text node is already inside a link
            if node.find_parent('a'):
                continue

            # we may have surrounding " | " etc; we rebuild by splitting the text node
            parts = re.split(rf'(\b{re.escape(label)}\b)', node)
            new_nodes = []
            for part in parts:
                if part == label:
                    a = soup.new_tag('a', href=url, target="_blank", rel="noopener noreferrer")
                    a.string = label
                    new_nodes.append(a)
                else:
                    if part:
                        new_nodes.append(NavigableString(part))

            # replace the original text node with the new sequence
            for i, nn in enumerate(new_nodes):
                if i == 0:
                    node.replace_with(nn)
                else:
                    # insert after the last inserted node
                    new_nodes[i-1].insert_after(nn)

    return str(soup)

def replace_placeholder_links(md_text: str, links: dict) -> str:
    """Turn [Label](), [Label](#), or bare [Label] into real links from LINK_MAP."""
    # [Label]() or [Label](#)
    md_text = re.sub(
        r'\[([^\]]+)\]\((?:#|)?\)',
        lambda m: f'[{m.group(1)}]({links.get(m.group(1).strip(), "#")})',
        md_text,
    )
    # bare [Label]
    md_text = re.sub(
        r'\[([^\]]+)\](?!\()',
        lambda m: f'[{m.group(1)}]({links.get(m.group(1).strip(), "#")})',
        md_text,
    )
    return md_text

def make_all_links_clickable(html_text: str) -> str:
    """Auto-link plain emails and bare URLs. Does NOT touch existing <a> tags."""
    # emails
    html_text = re.sub(
        r'(?<!href=")([\w\.-]+@[\w\.-]+\.\w+)',
        r'<a href="mailto:\1">\1</a>',
        html_text
    )
    # http(s) URLs
    html_text = re.sub(
        r'(?<!href=")(https?://[^\s"<]+)',
        r'<a href="\1">\1</a>',
        html_text
    )
    # www.* URLs
    html_text = re.sub(
        r'(?<!href=")(www\.[^\s"<]+)',
        r'<a href="https://\1">\1</a>',
        html_text
    )
    return html_text

def autolink_and_normalize(html_text: str) -> str:
    """
    - Convert plain emails / URLs / www.* to anchors
    - Ensure www.* becomes https://www.*
    - Add target/rel (harmless for PDF; useful for HTML preview)
    """
    # emails
    html_text = re.sub(
        r'(?<!href=")([\w\.-]+@[\w\.-]+\.\w+)',
        r'<a href="mailto:\1">\1</a>',
        html_text
    )
    # http(s)
    html_text = re.sub(
        r'(?<!href=")(https?://[^\s"<]+)',
        r'<a href="\1">\1</a>',
        html_text
    )
    # www.*
    html_text = re.sub(
        r'(?<!href=")(www\.[^\s"<]+)',
        r'<a href="https://\1">\1</a>',
        html_text
    )

    # normalize & annotate all anchors
    soup = BeautifulSoup(html_text, "html.parser")
    for a in soup.find_all('a', href=True):
        href = (a['href'] or '').strip()
        if href.startswith('www.'):
            a['href'] = f'https://{href}'
        a['target'] = '_blank'
        a['rel'] = 'noopener noreferrer'
    return str(soup)


@app.route('/contact', methods=['GET', 'POST'])
def contact():
   if request.method == 'POST':
       name = request.form.get('name')
       email = request.form.get('email')
       message = request.form.get('message')
       flash('Thank you for contacting us!', 'success')
       return redirect(url_for('contact'))
   return render_template('contact.html')


@app.route('/privacy-policy')
def privacy_policy():
   return render_template('privacy_policy.html')


@app.route('/terms')
def terms():
   return render_template('terms.html')


def remove_header_emails(text):
   email_pattern = re.compile(r'^[\s>]*[\w\.-]+@[\w\.-]+\.\w+[\s<]*$', re.IGNORECASE)
   lines = text.splitlines()
   new_lines = []
   for line in lines:
       if not new_lines and email_pattern.match(line.strip()):
           continue
       else:
           new_lines.append(line)
   return "\n".join(new_lines)


@app.route('/download_template_pdf_alt', methods=['POST'], endpoint='download_pdf_two')
@login_required
def download_template_pdf_alt():
   data = request.json
   content = data['content']
   html_ready = make_all_links_clickable(content)
   html_rendered = render_template('resume_pdf_template.html', content=html_ready)
   pdf = pdfkit.from_string(html_rendered, False, options={'enable-local-file-access': None})
   response = make_response(pdf)
   response.headers['Content-Type'] = 'application/pdf'
   response.headers['Content-Disposition'] = 'attachment; filename=resume.pdf'
   return response

def md_to_html_with_links(md_text):
    html = markdown.markdown(md_text, extensions=['extra'])
    html = bleach.linkify(
        html,
        callbacks=[bleach.callbacks.nofollow, bleach.callbacks.target_blank],
        parse_email=True
    )
    return html

@app.route('/download-resume', methods=['POST'])
def download_resume():
   return send_file(
       "resume.pdf",
       as_attachment=True,
       download_name="resume.pdf",
       mimetype="application/pdf"
   )


@app.route('/_debug_pdf_html', methods=['POST'])
@login_required
def _debug_pdf_html():
    data = request.get_json()
    content = replace_placeholder_links(data.get('content',''), LINK_MAP)
    html = markdown.markdown(content, extensions=['extra','sane_lists'])
    html = make_all_links_clickable(html)
    soup = BeautifulSoup(html, "html.parser")
    for a in soup.find_all('a', href=True):
        if a['href'].startswith('www.'):
            a['href'] = 'https://' + a['href']
    return str(soup)  # open this in browser and click links


@app.context_processor
def inject_now():
   return {'now': datetime.now}


@app.errorhandler(429)
def too_many_requests(e):
   return render_template("429.html"), 429


if __name__ == "__main__":
   with app.app_context():
       db.create_all()
   app.run(debug=True, port=5002)
